"""DOCX document ingestion using python-docx."""

from __future__ import annotations

from pathlib import Path

from ingestion.base import Document, Ingestor


class DocxIngestor(Ingestor):
    """Extract text, tables, and images from DOCX files."""

    def ingest(self, file_path: str | Path) -> Document:
        path = self._validate_path(file_path)
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("python-docx is required: pip install python-docx") from exc

        text_chunks: list[str] = []
        images: list[bytes] = []

        doc = DocxDocument(str(path))
        metadata = {
            "format": "docx",
            "filename": path.name,
        }

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_chunks.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                text_chunks.append("\n".join(rows))

        # Extract images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    images.append(rel.target_part.blob)
                except Exception:
                    continue

        return Document(text_chunks=text_chunks, images=images, metadata=metadata)
